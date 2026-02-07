# -*- coding: utf-8 -*-
"""智能工具人机交互与远程操控系统 - 项目立项文档生成器"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime
import os

def set_font(run, font_name='宋体', size=12, bold=False):
    """设置字体"""
    run.font.name = font_name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if bold:
        run.bold = True

def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        set_font(run, '黑体', 16 if level == 1 else 14, True)
    return heading

def add_para(doc, text, bold=False, indent=0):
    """添加段落"""
    p = doc.add_paragraph()
    if indent > 0:
        p.paragraph_format.left_indent = Inches(indent * 0.3)
    run = p.add_run(text)
    set_font(run, '宋体', 12, bold)
    p.paragraph_format.line_spacing = 1.5
    return p

def create_table(doc, headers, rows):
    """创建表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                set_font(run, '黑体', 11, True)
    
    # 数据行
    for row_data in rows:
        cells = table.add_row().cells
        for i, data in enumerate(row_data):
            cells[i].text = str(data)
            for para in cells[i].paragraphs:
                for run in para.runs:
                    set_font(run, '宋体', 10)
    return table

def main():
    doc = Document()
    
    # 设置默认样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)
    
    # ==================== 封面 ====================
    title = doc.add_heading('智能工具人机交互与远程操控系统', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_font(run, '黑体', 22, True)
    
    subtitle = doc.add_heading('项目立项申报书', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        set_font(run, '黑体', 18, True)
    
    doc.add_paragraph('\n' * 3)
    
    # 基本信息
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_text = (
        f'项目名称：智能工具人机交互与远程操控系统\n'
        f'项目版本：V1.0\n'
        f'申报日期：2025年11月08日\n'
        f'所属领域：智能工具人机交互 / 企业级系统管理\n'
        f'项目类别：软件开发 / 大数据应用\n'
    )
    run = info.add_run(info_text)
    set_font(run, '楷体', 14)
    
    doc.add_page_break()
    
    # ==================== 目录 ====================
    add_heading_custom(doc, '目 录', 1)
    toc = [
        '一、项目概述',
        '二、项目背景与必要性',
        '三、国内外研究现状与技术发展趋势',
        '四、项目目标与主要内容',
        '五、技术方案与实施路线',
        '六、关键技术与创新点',
        '七、系统架构与功能模块',
        '八、技术指标与性能要求',
        '九、项目实施计划与进度安排',
        '十、项目预算与资源配置',
        '十一、风险分析与应对措施',
        '十二、预期成果与社会经济效益',
        '十三、知识产权与成果保护',
        '十四、项目团队与保障条件',
        '十五、总结与展望',
        '十六、项目成功关键因素',
    ]
    for item in toc:
        add_para(doc, item, indent=1)
    
    doc.add_page_break()
    
    # ==================== 一、项目概述 ====================
    add_heading_custom(doc, '一、项目概述', 1)
    
    add_para(doc,
        '智能工具人机交互与远程操控系统（Smart Control System）是一套拟开发的、面向企业级应用的综合性管控平台。'
        '项目旨在针对现代企业在设备管理和远程操控领域面临的痛点，打造一个高效、安全、智能的解决方案。'
        '本系统计划基于Go语言（1.20+版本）进行开发，采用分布式架构设计，并整合Gin Web框架、SQLite数据库、'
        'WebSocket实时通信、noVNC远程桌面等先进技术，以实现对企业智能设备的全方位、精细化管控。')
    
    add_para(doc,
        '项目规划建设10大核心功能模块，涵盖系统监控、语音交互、远程控制、视频监控、硬件检测、'
        '软件更新、数据同步、异常报警、性能分析和维护日志等关键业务场景。通过实时数据采集、智能分析决策和可视化展示，'
        '项目目标是帮助企业显著提升运维效率、降低人力成本、保障系统稳定运行。项目成功实施后，'
        '预计可为用户实现运维效率提升60%以上，人力成本降低40%以上，系统可用性达到99.9%以上的应用效果。')
    
    add_heading_custom(doc, '1.1 项目基本信息', 2)
    
    basic_info = [
        ['软件全称', '智能工具人机交互与远程操控系统'],
        ['英文名称', 'Smart Control System'],
        ['版本号', 'V1.0'],
        ['开发语言', 'Go 1.20+'],
        ['源程序量', '约 20,000 行'],
        ['开发环境', 'LiteIDE / VS Code， Windows / Linux'],
        ['运行平台', 'Go 1.20 + Alpine Linux / Windows / macOS'],
        ['核心技术', 'Gin + SQLite + WebSocket + noVNC'],
        ['应用领域', '智能工具人机交互 / 企业级系统管理'],
        ['目标用户', '企业运维团队、系统管理员、设备管理人员'],
    ]
    create_table(doc, ['项目', '内容'], basic_info)
    
    add_heading_custom(doc, '1.2 系统核心特性', 2)
    
    features = [
        ('企业级安全认证', 
         '计划采用bcrypt密码加密算法，保障用户密码安全；实现Token持久化会话管理，24小时自动过期；'
         '支持API Token认证保护，防止未授权访问；提供CORS跨域支持和SQL注入防护。'),
        
        ('高性能架构', 
         '充分利用Go语言的高并发特性，支持大规模并发请求；通过数据库索引优化，目标查询速度提升50-80%；'
         '所有列表接口支持分页功能；实现请求限流（目标100 req/min），防止系统过载。'),
        
        ('实时监控告警', 
         '基于WebSocket技术实现CPU、内存、磁盘、网络指标的实时推送；支持阈值自定义配置；'
         '智能告警系统目标在10秒内检测异常并推送告警；提供多级告警机制和确认响应功能。'),
        
        ('模块化设计', 
         '10大独立功能模块，各模块职责清晰；支持灵活扩展和定制化开发；'
         '采用RESTful API设计，便于系统集成；提供完善的API文档和接口说明。'),
        
        ('跨平台部署', 
         '支持Windows、Linux、macOS等多种操作系统；提供Docker容器化部署方案；'
         '支持systemd服务管理；可编译为单一可执行文件，无需额外依赖。'),
        
        ('现代化界面', 
         '采用响应式Web设计，支持多终端访问，提供直观的数据可视化图表，'
         '目标为用户提供低学习成本的操作体验。'),
    ]
    
    for i, (title, desc) in enumerate(features, 1):
        add_para(doc, f'{i}. {title}', True, 1)
        add_para(doc, desc, indent=2)
    
    doc.add_page_break()
    
    # ==================== 二、项目背景与必要性 ====================
    print("正在生成：二、项目背景与必要性...")
    add_heading_custom(doc, '二、项目背景与必要性', 1)
    
    add_heading_custom(doc, '2.1 行业背景', 2)
    add_para(doc,
        '随着工业4.0和智能制造战略的深入推进，企业对智能设备的管理需求日益复杂化。'
        '根据工信部数据，2023年我国智能制造装备产业规模超过3万亿元，'
        '企业数字化转型进入深水区，对智能化管控系统的需求呈现爆发式增长。')
    
    add_para(doc, '当前企业面临的核心痛点：')
    pain_points = [
        '设备分散难管理：多地分布式设备缺乏统一管控平台',
        '实时监控能力弱：故障发现滞后，无法及时响应',
        '远程操作不便：传统工具安全性差、功能单一',
        '数据孤岛严重：各系统独立运行，数据无法互通',
        '告警响应迟缓：缺少智能告警，问题发现依赖人工',
        '操作审计缺失：维护操作无法追溯，存在安全隐患',
    ]
    for i, point in enumerate(pain_points, 1):
        add_para(doc, f'({i}) {point}', indent=1)
    
    add_heading_custom(doc, '2.2 项目必要性', 2)
    necessity = [
        ('提升运维效率', '可提升60%以上，降低人力成本40%以上'),
        ('保障业务连续性', '系统可用性提升至99.9%以上'),
        ('降低安全风险', '符合等保2.0和行业合规要求'),
        ('支撑数字化转型', '打破数据孤岛，推动智能化转型'),
        ('赋能管理决策', '提供多维度指标，支撑科学决策'),
    ]
    for i, (title, content) in enumerate(necessity, 1):
        add_para(doc, f'{i}. {title}', True, 1)
        add_para(doc, content, indent=2)
    
    add_heading_custom(doc, '2.3 市场需求分析', 2)
    market_data = [
        ['制造业', '智能工厂设备管理', '规上企业超10万家', '高'],
        ['能源行业', '电力/石油设施监控', '数千个站点', '高'],
        ['物流行业', '仓储设备管理', '超5万个物流中心', '中'],
        ['医疗行业', '医疗设备维护', '超3万家医疗机构', '中'],
        ['教育行业', '机房设备管理', '数万所学校', '中'],
        ['互联网/IDC', '服务器集群管理', '数千家数据中心', '高'],
    ]
    create_table(doc, ['行业', '应用场景', '市场规模', '需求程度'], market_data)
    add_para(doc,
        '\n预计到2025年国内市场规模将突破500亿元，本项目具有广阔的市场空间和商业化前景。')
    
    doc.add_page_break()
    
    # ==================== 三、国内外研究现状 ====================
    print("正在生成：三、国内外研究现状...")
    add_heading_custom(doc, '三、国内外研究现状与技术发展趋势', 1)
    
    add_heading_custom(doc, '3.1 国外研究现状', 2)
    international = [
        ('TeamViewer', '全球领先，但价格昂贵，数据存储海外'),
        ('Nagios/Zabbix', '开源系统，配置复杂，学习曲线陡峭'),
        ('VMware vCenter', '虚拟化管理，授权费用高昂'),
        ('AWS/Azure IoT', '云平台服务，依赖云服务，成本不可控'),
    ]
    for i, (name, desc) in enumerate(international, 1):
        add_para(doc, f'{i}. {name}：{desc}', indent=1)
    
    add_heading_custom(doc, '3.2 国内研究现状', 2)
    domestic = [
        ('向日葵/ToDesk', '界面友好，但功能单一，缺少企业级特性'),
        ('云厂商IoT平台', '技术先进，但私有化部署成本高'),
        ('自研系统', '针对性强，但技术老旧，维护成本高'),
    ]
    for i, (name, desc) in enumerate(domestic, 1):
        add_para(doc, f'{i}. {name}：{desc}', indent=1)
    
    add_heading_custom(doc, '3.3 技术发展趋势', 2)
    trends = [
        '云边协同：边缘计算与云计算融合',
        'AI赋能：故障预测、智能调度',
        '微服务化：提升系统弹性和可维护性',
        '容器化部署：简化部署和运维',
        '零信任安全：强化身份认证和访问控制',
        '数字孪生：支持仿真和预演',
    ]
    for i, item in enumerate(trends, 1):
        add_para(doc, f'{i}. {item}', indent=1)
    
    doc.add_page_break()
    
    # 调用补充内容生成模块
    print("正在生成剩余章节...")
    from proposal_content import generate_remaining_sections
    generate_remaining_sections(doc, add_heading_custom, add_para, create_table)
    
    # 保存文档 - 固定文件名，覆盖旧文件
    base_filename = '智能工具人机交互与远程操控系统_项目立项申报书'
    output_file = f'{base_filename}.docx'
    
    try:
        doc.save(output_file)
        print('\n[SUCCESS] 文档生成成功（已覆盖同名文件）！')
        print(f'文件位置：{os.path.abspath(output_file)}')
        print('文档包含16个完整章节，内容详尽，格式规范')
        print('建议使用Microsoft Word或WPS打开查看')
    except PermissionError:
        # 如果仍然失败，尝试保存到临时位置
        temp_file = f'{base_filename}_temp_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx'
        try:
            doc.save(temp_file)
            print('\n[WARNING] 原文件被占用，已保存到临时文件！')
            print(f'文件位置：{os.path.abspath(temp_file)}')
            print('\n提示：请关闭已打开的 Word 文档后重新运行')
        except Exception as e:
            print(f'\n[ERROR] 文件保存失败：{e}')
            print('请确保：')
            print('1. 关闭所有打开的相关 Word 文档')
            print('2. 检查目录是否有写入权限')
            print('3. 磁盘空间是否充足')
            raise
    print(f'\n文档包含以下章节：')
    print('  一、项目概述')
    print('  二、项目背景与必要性')
    print('  三、国内外研究现状与技术发展趋势')
    print('  四、项目目标与主要内容')
    print('  五、技术方案与实施路线')
    print('  六、关键技术与创新点')
    print('  七、系统架构与功能模块')
    print('  八、技术指标与性能要求')
    print('  九、项目实施计划与进度安排')
    print('  十、项目预算与资源配置')
    print('  十一、风险分析与应对措施')
    print('  十二、预期成果与社会经济效益')
    print('  十三、知识产权与成果保护')
    print('  十四、项目团队与保障条件')
    print('  十五、总结与展望')
    print('  十六、项目成功关键因素')

if __name__ == '__main__':
    print('='*60)
    print('智能工具人机交互与远程操控系统')
    print('项目立项申报书 - 文档生成器')
    print('='*60)
    print('\n开始生成文档...\n')
    main()
    print('\n'+'='*60)
